from __future__ import annotations

import hashlib
from datetime import datetime
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.tools.base import Tool, ToolContext, ToolResult


class DocxGeneratorTool(Tool):
    name = "doc.generate_docx"
    risk_tier = 1
    description = "Generate a DOCX document and store it as an artifact."

    def call(self, ctx: ToolContext, **kwargs: Any) -> ToolResult:
        title = kwargs.get("title", "Document")
        paragraphs = kwargs.get("paragraphs", [])
        filename = kwargs.get("filename", f"{title.replace(' ', '_')}.docx")

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
        doc.add_paragraph("")
        for p in paragraphs:
            doc.add_paragraph(str(p))

        path = ctx.evidence.root / "artifacts" / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))

        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        eid = ctx.evidence.next_evidence_id("DOCX")
        ctx.evidence.add_artifact(evidence_id=eid, type="docx", path=path, sha256=digest, metadata={"title": title})

        return ToolResult(ok=True, output={"filename": filename, "path": str(path), "sha256": digest, "artifacts": [{"evidence_id": eid, "type": "docx", "path": str(path), "sha256": digest}]}, evidence_ids=[eid])


class PdfGeneratorTool(Tool):
    name = "doc.generate_pdf"
    risk_tier = 1
    description = "Generate a PDF document and store it as an artifact."

    def call(self, ctx: ToolContext, **kwargs: Any) -> ToolResult:
        title = kwargs.get("title", "Report")
        lines = kwargs.get("lines", [])
        filename = kwargs.get("filename", f"{title.replace(' ', '_')}.pdf")

        path = ctx.evidence.root / "artifacts" / filename
        path.parent.mkdir(parents=True, exist_ok=True)

        c = canvas.Canvas(str(path), pagesize=letter)
        width, height = letter
        y = height - 72
        c.setTitle(title)
        c.drawString(72, y, title)
        y -= 24
        c.drawString(72, y, f"Generated: {datetime.utcnow().isoformat()}Z")
        y -= 36
        for line in lines:
            if y < 72:
                c.showPage()
                y = height - 72
            c.drawString(72, y, str(line)[:110])
            y -= 16
        c.save()

        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        eid = ctx.evidence.next_evidence_id("PDF")
        ctx.evidence.add_artifact(evidence_id=eid, type="pdf", path=path, sha256=digest, metadata={"title": title})

        return ToolResult(ok=True, output={"filename": filename, "path": str(path), "sha256": digest, "artifacts": [{"evidence_id": eid, "type": "pdf", "path": str(path), "sha256": digest}]}, evidence_ids=[eid])
